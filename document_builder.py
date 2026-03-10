"""
Word document builder for study notes
"""

from docx import Document
from docx.shared import Inches, Cm, Pt
from docx.enum.text import WD_COLOR_INDEX
from pathlib import Path
import datetime
from config import SECTION_DURATION, MAX_IMAGE_WIDTH_CM

def build_document(video_info, segments, approved_frames, output_name, temp_dir):
    """
    Builds the final Word document.

    Returns: path to output file
    """
    doc = Document()

    # Cover page
    title = doc.add_heading('LMS Video Study Notes', 0)
    doc.add_paragraph(f"Video Source: {video_info['platform'].title()}")
    doc.add_paragraph(f"URL: {video_info['video_url']}")
    doc.add_paragraph(f"Generated: {datetime.datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
    doc.add_paragraph(f"Transcript Segments: {len(segments)}")
    doc.add_paragraph(f"Screenshots: {len(approved_frames)}")
    doc.add_page_break()

    # Group frames by 2-minute sections
    frame_groups = {}
    for frame in approved_frames:
        section_start = int(frame['timestamp_sec'] // SECTION_DURATION) * SECTION_DURATION
        section_key = f"{seconds_to_hms(section_start)} – {seconds_to_hms(section_start + SECTION_DURATION)}"
        if section_key not in frame_groups:
            frame_groups[section_key] = []
        frame_groups[section_key].append(frame)

    # Sort sections
    sorted_sections = sorted(frame_groups.keys(), key=lambda x: hms_to_seconds(x.split(' – ')[0]))

    for section_title in sorted_sections:
        # Section heading
        doc.add_heading(f"▶ {section_title}", level=1)

        # Transcript for this section
        section_start = hms_to_seconds(section_title.split(' – ')[0])
        section_end = section_start + SECTION_DURATION
        transcript_text = get_section_transcript(segments, section_start, section_end)
        para = doc.add_paragraph()
        for segment in transcript_text:
            # Add timestamp in gray italic
            run = para.add_run(f"[{seconds_to_hms(segment['start'])}] ")
            run.font.color.rgb = None  # Gray
            run.font.italic = True
            run.font.size = Pt(10)
            # Add text
            para.add_run(segment['text'] + " ")

        # Screenshots
        for frame in frame_groups[section_title]:
            doc.add_picture(frame['frame_path'], width=Cm(MAX_IMAGE_WIDTH_CM))
            caption = f"📸 {frame.get('category', 'other').title()} · {frame['timestamp_str']} · {frame.get('reason', '')}"
            para = doc.add_paragraph(caption)
            para.style.font.size = Pt(10)
            para.style.font.color.rgb = None  # Blue

        doc.add_page_break()

    # Footer with page numbers
    for section in doc.sections:
        footer = section.footer
        para = footer.paragraphs[0]
        para.text = "Page {PAGE} of {NUMPAGES}"

    output_path = Path.cwd() / f"{output_name}.docx"
    doc.save(str(output_path))
    return str(output_path)

def get_section_transcript(segments, start_sec, end_sec):
    """Get transcript segments for a time section."""
    return [s for s in segments if start_sec <= s['start'] < end_sec]

def seconds_to_hms(seconds):
    """Convert seconds to HH:MM:SS."""
    h = int(seconds // 3600)
    m = int((seconds % 3600) // 60)
    s = int(seconds % 60)
    return f"{h:02d}:{m:02d}:{s:02d}"

def hms_to_seconds(hms):
    """Convert HH:MM:SS to seconds."""
    h, m, s = map(int, hms.split(':'))
    return h * 3600 + m * 60 + s